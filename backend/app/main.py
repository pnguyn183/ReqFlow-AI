from __future__ import annotations

import base64
import csv
import hashlib
import hmac
import io
import json
import os
import re
import secrets
import time
from datetime import datetime, timedelta, timezone
from typing import Any

from docx import Document as DocxDocument
from fastapi import Depends, FastAPI, File, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from prometheus_client import CONTENT_TYPE_LATEST, Counter, Histogram, generate_latest
from pydantic import BaseModel, Field
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from sqlalchemy import Boolean, DateTime, Float, ForeignKey, Integer, String, Text, create_engine, func, select
from sqlalchemy.orm import DeclarativeBase, Mapped, Session, mapped_column, sessionmaker

try:
    from openai import OpenAI
except ImportError:  # pragma: no cover - local fallback remains available
    OpenAI = None


def utcnow() -> datetime:
    return datetime.now(timezone.utc)


DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./reqflow.db")
SECRET_KEY = os.getenv("SECRET_KEY", "reqflow-local-secret-change-in-production")
TOKEN_TTL_MINUTES = int(os.getenv("TOKEN_TTL_MINUTES", "480"))
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.4-mini")

connect_args = {"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}
engine = create_engine(DATABASE_URL, connect_args=connect_args, pool_pre_ping=True)
SessionLocal = sessionmaker(bind=engine, expire_on_commit=False)


class Base(DeclarativeBase):
    pass


class User(Base):
    __tablename__ = "users"

    id: Mapped[int] = mapped_column(primary_key=True)
    username: Mapped[str] = mapped_column(String(80), unique=True, index=True)
    full_name: Mapped[str] = mapped_column(String(160))
    role: Mapped[str] = mapped_column(String(40), default="business_analyst")
    password_hash: Mapped[str] = mapped_column(String(200))
    active: Mapped[bool] = mapped_column(Boolean, default=True)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)


class Requirement(Base):
    __tablename__ = "requirements"

    id: Mapped[int] = mapped_column(primary_key=True)
    code: Mapped[str] = mapped_column(String(30), unique=True, index=True)
    title: Mapped[str] = mapped_column(String(240))
    description: Mapped[str] = mapped_column(Text)
    requirement_type: Mapped[str] = mapped_column(String(40), default="functional")
    category: Mapped[str] = mapped_column(String(80), default="General")
    priority: Mapped[str] = mapped_column(String(20), default="Should")
    status: Mapped[str] = mapped_column(String(30), default="Draft")
    owner: Mapped[str] = mapped_column(String(120), default="Unassigned")
    project: Mapped[str] = mapped_column(String(120), default="ReqFlow AI")
    source: Mapped[str] = mapped_column(String(80), default="Workshop")
    business_value: Mapped[float] = mapped_column(Float, default=5)
    complexity: Mapped[float] = mapped_column(Float, default=3)
    reach: Mapped[float] = mapped_column(Float, default=100)
    impact: Mapped[float] = mapped_column(Float, default=2)
    confidence: Mapped[float] = mapped_column(Float, default=80)
    effort: Mapped[float] = mapped_column(Float, default=3)
    rice_score: Mapped[float] = mapped_column(Float, default=0)
    version: Mapped[int] = mapped_column(Integer, default=1)
    approved_by: Mapped[str | None] = mapped_column(String(120), nullable=True)
    created_by: Mapped[str] = mapped_column(String(120))
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)
    updated_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow, onupdate=utcnow)


class RequirementVersion(Base):
    __tablename__ = "requirement_versions"

    id: Mapped[int] = mapped_column(primary_key=True)
    requirement_id: Mapped[int] = mapped_column(ForeignKey("requirements.id"), index=True)
    version_number: Mapped[int] = mapped_column(Integer)
    snapshot_json: Mapped[str] = mapped_column(Text)
    change_summary: Mapped[str] = mapped_column(String(300))
    created_by: Mapped[str] = mapped_column(String(120))
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)


class UserStory(Base):
    __tablename__ = "user_stories"

    id: Mapped[int] = mapped_column(primary_key=True)
    requirement_id: Mapped[int] = mapped_column(ForeignKey("requirements.id"), index=True)
    title: Mapped[str] = mapped_column(String(240))
    persona: Mapped[str] = mapped_column(String(120))
    goal: Mapped[str] = mapped_column(Text)
    benefit: Mapped[str] = mapped_column(Text)
    acceptance_criteria_json: Mapped[str] = mapped_column(Text)
    definition_of_done_json: Mapped[str] = mapped_column(Text)
    status: Mapped[str] = mapped_column(String(30), default="Ready")
    story_points: Mapped[int] = mapped_column(Integer, default=3)
    sprint: Mapped[str] = mapped_column(String(40), default="Backlog")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)


class TraceLink(Base):
    __tablename__ = "trace_links"

    id: Mapped[int] = mapped_column(primary_key=True)
    business_goal: Mapped[str] = mapped_column(String(240))
    requirement_id: Mapped[int] = mapped_column(ForeignKey("requirements.id"), index=True)
    user_story_id: Mapped[int | None] = mapped_column(ForeignKey("user_stories.id"), nullable=True)
    task: Mapped[str] = mapped_column(String(240), default="Not linked")
    test_case: Mapped[str] = mapped_column(String(240), default="Not linked")
    coverage_status: Mapped[str] = mapped_column(String(30), default="Partial")


class ChangeRequest(Base):
    __tablename__ = "change_requests"

    id: Mapped[int] = mapped_column(primary_key=True)
    requirement_id: Mapped[int] = mapped_column(ForeignKey("requirements.id"), index=True)
    summary: Mapped[str] = mapped_column(String(300))
    reason: Mapped[str] = mapped_column(Text)
    impact: Mapped[str] = mapped_column(String(30), default="Medium")
    status: Mapped[str] = mapped_column(String(30), default="Pending")
    requested_by: Mapped[str] = mapped_column(String(120))
    reviewed_by: Mapped[str | None] = mapped_column(String(120), nullable=True)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)
    reviewed_at: Mapped[datetime | None] = mapped_column(DateTime(timezone=True), nullable=True)


class GeneratedDocument(Base):
    __tablename__ = "generated_documents"

    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str] = mapped_column(String(240))
    document_type: Mapped[str] = mapped_column(String(30))
    content: Mapped[str] = mapped_column(Text)
    status: Mapped[str] = mapped_column(String(30), default="Generated")
    requirement_count: Mapped[int] = mapped_column(Integer, default=0)
    generated_by: Mapped[str] = mapped_column(String(120))
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)


class AuditLog(Base):
    __tablename__ = "audit_logs"

    id: Mapped[int] = mapped_column(primary_key=True)
    actor: Mapped[str] = mapped_column(String(120))
    action: Mapped[str] = mapped_column(String(100))
    entity_type: Mapped[str] = mapped_column(String(80))
    entity_id: Mapped[str] = mapped_column(String(80))
    detail: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)


class LoginPayload(BaseModel):
    username: str
    password: str


class RequirementCreate(BaseModel):
    title: str = Field(min_length=3, max_length=240)
    description: str = Field(min_length=10)
    requirement_type: str = "functional"
    category: str = "General"
    priority: str = "Should"
    status: str = "Draft"
    owner: str = "Unassigned"
    project: str = "ReqFlow AI"
    source: str = "Workshop"
    business_value: float = Field(default=5, ge=1, le=10)
    complexity: float = Field(default=3, ge=1, le=10)
    reach: float = Field(default=100, ge=1)
    impact: float = Field(default=2, ge=0.25, le=5)
    confidence: float = Field(default=80, ge=1, le=100)
    effort: float = Field(default=3, ge=0.5)


class RequirementUpdate(BaseModel):
    title: str | None = None
    description: str | None = None
    requirement_type: str | None = None
    category: str | None = None
    priority: str | None = None
    status: str | None = None
    owner: str | None = None
    business_value: float | None = None
    complexity: float | None = None
    change_summary: str = "Requirement updated"


class DocumentGeneratePayload(BaseModel):
    document_type: str = Field(pattern="^(BRD|SRS|BACKLOG)$")
    name: str | None = None
    requirement_ids: list[int] = []


class ProcessGeneratePayload(BaseModel):
    model_type: str = Field(pattern="^(USE_CASE|BPMN|ERD)$")
    requirement_id: int | None = None


class TraceCreatePayload(BaseModel):
    business_goal: str
    requirement_id: int
    user_story_id: int | None = None
    task: str = "Not linked"
    test_case: str = "Not linked"
    coverage_status: str = "Partial"


class ChangeCreatePayload(BaseModel):
    requirement_id: int
    summary: str = Field(min_length=5)
    reason: str = Field(min_length=5)
    impact: str = "Medium"


class ChangeDecisionPayload(BaseModel):
    decision: str = Field(pattern="^(Approved|Rejected)$")


class AssistantPayload(BaseModel):
    question: str = Field(min_length=3)
    requirement_id: int | None = None


ROLE_PERMISSIONS = {
    "admin": {"read", "write", "approve", "manage_users", "export"},
    "business_analyst": {"read", "write", "export"},
    "product_owner": {"read", "write", "approve", "export"},
    "project_manager": {"read", "approve", "export"},
    "viewer": {"read"},
}


def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 120_000).hex()
    return f"{salt}${digest}"


def verify_password(password: str, stored: str) -> bool:
    salt, expected = stored.split("$", 1)
    actual = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 120_000).hex()
    return hmac.compare_digest(actual, expected)


def encode_token(user: User) -> str:
    payload = {
        "sub": user.username,
        "role": user.role,
        "exp": int((utcnow() + timedelta(minutes=TOKEN_TTL_MINUTES)).timestamp()),
    }
    raw = base64.urlsafe_b64encode(json.dumps(payload).encode()).decode().rstrip("=")
    signature = hmac.new(SECRET_KEY.encode(), raw.encode(), hashlib.sha256).hexdigest()
    return f"{raw}.{signature}"


def decode_token(token: str) -> dict[str, Any]:
    try:
        raw, signature = token.split(".", 1)
        expected = hmac.new(SECRET_KEY.encode(), raw.encode(), hashlib.sha256).hexdigest()
        if not hmac.compare_digest(signature, expected):
            raise ValueError("Invalid signature")
        padded = raw + "=" * (-len(raw) % 4)
        payload = json.loads(base64.urlsafe_b64decode(padded).decode())
        if payload["exp"] < int(utcnow().timestamp()):
            raise ValueError("Expired token")
        return payload
    except (ValueError, KeyError, json.JSONDecodeError) as exc:
        raise HTTPException(status_code=401, detail="Invalid or expired token") from exc


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def get_current_user(request: Request, db: Session = Depends(get_db)) -> User:
    header = request.headers.get("Authorization", "")
    if not header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Authentication required")
    payload = decode_token(header.removeprefix("Bearer ").strip())
    user = db.scalar(select(User).where(User.username == payload["sub"], User.active.is_(True)))
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


def require(permission: str):
    def dependency(user: User = Depends(get_current_user)) -> User:
        if permission not in ROLE_PERMISSIONS.get(user.role, set()):
            raise HTTPException(status_code=403, detail=f"Missing permission: {permission}")
        return user

    return dependency


def requirement_snapshot(requirement: Requirement) -> dict[str, Any]:
    return {
        "code": requirement.code,
        "title": requirement.title,
        "description": requirement.description,
        "requirement_type": requirement.requirement_type,
        "category": requirement.category,
        "priority": requirement.priority,
        "status": requirement.status,
        "owner": requirement.owner,
        "version": requirement.version,
    }


def requirement_dict(requirement: Requirement) -> dict[str, Any]:
    return {
        **requirement_snapshot(requirement),
        "id": requirement.id,
        "project": requirement.project,
        "source": requirement.source,
        "business_value": requirement.business_value,
        "complexity": requirement.complexity,
        "reach": requirement.reach,
        "impact": requirement.impact,
        "confidence": requirement.confidence,
        "effort": requirement.effort,
        "rice_score": round(requirement.rice_score, 1),
        "approved_by": requirement.approved_by,
        "created_by": requirement.created_by,
        "created_at": requirement.created_at.isoformat(),
        "updated_at": requirement.updated_at.isoformat(),
    }


def story_dict(story: UserStory, requirement_code: str = "") -> dict[str, Any]:
    return {
        "id": story.id,
        "requirement_id": story.requirement_id,
        "requirement_code": requirement_code,
        "title": story.title,
        "persona": story.persona,
        "goal": story.goal,
        "benefit": story.benefit,
        "acceptance_criteria": json.loads(story.acceptance_criteria_json),
        "definition_of_done": json.loads(story.definition_of_done_json),
        "status": story.status,
        "story_points": story.story_points,
        "sprint": story.sprint,
    }


def audit(db: Session, user: User | str, action: str, entity_type: str, entity_id: Any, detail: str = "") -> None:
    actor = user.username if isinstance(user, User) else user
    db.add(AuditLog(actor=actor, action=action, entity_type=entity_type, entity_id=str(entity_id), detail=detail))


def next_requirement_code(db: Session) -> str:
    highest = db.scalar(select(func.max(Requirement.id))) or 0
    return f"REQ-{highest + 1:04d}"


def calculate_rice(requirement: Requirement) -> float:
    return (requirement.reach * requirement.impact * (requirement.confidence / 100)) / max(requirement.effort, 0.5)


def local_analysis(text: str) -> dict[str, Any]:
    lowered = text.lower()
    nfr_map = {
        "Performance": ["response", "second", "latency", "fast", "performance"],
        "Security": ["security", "encrypt", "permission", "role", "authentication", "rbac"],
        "Availability": ["availability", "uptime", "99.9", "recover"],
        "Scalability": ["scale", "concurrent", "10,000", "volume"],
        "Auditability": ["audit", "history", "trace"],
    }
    nfr = [name for name, keywords in nfr_map.items() if any(keyword in lowered for keyword in keywords)]
    stakeholders = []
    stakeholder_map = {
        "Business Analyst": ["analyst", "requirement", "brd", "srs"],
        "Product Owner": ["backlog", "priority", "product owner"],
        "Project Manager": ["project", "sprint", "delivery"],
        "Administrator": ["admin", "role", "permission"],
        "End User": ["user", "customer", "submit"],
    }
    for name, keywords in stakeholder_map.items():
        if any(keyword in lowered for keyword in keywords):
            stakeholders.append(name)
    if not stakeholders:
        stakeholders = ["Business Analyst", "Product Owner"]
    ambiguous = []
    if not re.search(r"\b(must|shall|should|can|will)\b", lowered):
        ambiguous.append("The expected system behavior is not expressed with a testable modal verb.")
    if not re.search(r"\b(user|analyst|owner|manager|administrator|system)\b", lowered):
        ambiguous.append("The primary actor is not explicit.")
    if not re.search(r"\b(within|less than|more than|at least|percent|%)\b", lowered):
        ambiguous.append("No measurable acceptance threshold is stated.")
    quality = max(48, 96 - len(ambiguous) * 13)
    category = "Security" if "Security" in nfr else "Workflow" if "approval" in lowered else "Requirements"
    return {
        "classification": "non-functional" if nfr else "functional",
        "category": category,
        "functional_requirements": [
            sentence.strip() for sentence in re.split(r"[.!?]+", text) if len(sentence.strip()) > 18
        ][:4] or [text],
        "non_functional_requirements": nfr,
        "stakeholders": stakeholders,
        "ambiguities": ambiguous,
        "missing_requirements": [
            "Define failure and recovery behavior.",
            "Confirm authorization rules for each action.",
            "Add measurable acceptance criteria.",
        ][: max(1, len(ambiguous))],
        "quality_score": quality,
        "provider": "local",
    }


def openai_answer(instructions: str, prompt: str) -> str | None:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or OpenAI is None:
        return None
    try:
        client = OpenAI(api_key=api_key)
        response = client.responses.create(model=OPENAI_MODEL, instructions=instructions, input=prompt)
        return response.output_text
    except Exception:
        return None


def generate_story(requirement: Requirement) -> dict[str, Any]:
    persona_map = {
        "Security": "system administrator",
        "Workflow": "product owner",
        "Reporting": "project manager",
        "Requirements": "business analyst",
    }
    persona = persona_map.get(requirement.category, "business user")
    return {
        "title": requirement.title,
        "persona": persona,
        "goal": requirement.description.rstrip("."),
        "benefit": "the team can deliver a validated outcome with clear ownership",
        "acceptance_criteria": [
            f"Given an authenticated {persona}, when the action is completed, then the result is stored successfully.",
            "Validation errors are shown without losing submitted data.",
            "The action is recorded in the audit log.",
        ],
        "definition_of_done": [
            "Acceptance criteria pass in automated tests.",
            "RBAC rules are verified.",
            "Product Owner accepts the behavior.",
            "Technical and user documentation are updated.",
        ],
    }


def seed_database() -> None:
    Base.metadata.create_all(bind=engine)
    with SessionLocal() as db:
        if not db.scalar(select(func.count(User.id))):
            db.add_all(
                [
                    User(username="admin", full_name="System Administrator", role="admin", password_hash=hash_password("admin123")),
                    User(username="analyst", full_name="Hoang Phuc Nguyen", role="business_analyst", password_hash=hash_password("demo123")),
                    User(username="owner", full_name="Linh Tran", role="product_owner", password_hash=hash_password("demo123")),
                ]
            )
            db.commit()
        if db.scalar(select(func.count(Requirement.id))):
            return
        samples = [
            ("AI requirement analysis", "The system shall analyze submitted requirement text and extract functional requirements, non-functional requirements, stakeholders, and ambiguities.", "functional", "Requirements", "Must", "Approved", 10, 5),
            ("Requirement approval workflow", "A Product Owner must review and approve requirement changes before they become the active version.", "functional", "Workflow", "Must", "In Review", 9, 4),
            ("Document generation", "Business Analysts shall generate BRD and SRS documents from approved requirements and export them as PDF or DOCX.", "functional", "Reporting", "Must", "Approved", 9, 6),
            ("Traceability coverage", "The platform shall maintain links from business goals to requirements, user stories, delivery tasks, and test cases.", "functional", "Requirements", "Must", "Approved", 10, 4),
            ("Response time", "API responses should complete within 500 milliseconds for 95 percent of requests.", "non-functional", "Performance", "Should", "Draft", 7, 5),
            ("Role-based access", "Administrators shall assign roles and permissions while all protected actions are captured in an audit log.", "non-functional", "Security", "Must", "Approved", 10, 5),
            ("Backlog prioritization", "Product Owners can rank requirements using MoSCoW, RICE, and value versus complexity.", "functional", "Workflow", "Should", "In Review", 8, 3),
            ("Missing requirement suggestions", "The AI assistant should identify gaps, conflicting statements, and missing acceptance criteria.", "functional", "Requirements", "Could", "Draft", 7, 4),
        ]
        requirements: list[Requirement] = []
        for index, sample in enumerate(samples, start=1):
            title, description, req_type, category, priority, status, value, complexity = sample
            requirement = Requirement(
                code=f"REQ-{index:04d}", title=title, description=description, requirement_type=req_type,
                category=category, priority=priority, status=status, owner="Hoang Phuc Nguyen",
                business_value=value, complexity=complexity, reach=120 + index * 25, impact=2.5,
                confidence=85, effort=max(2, complexity), created_by="analyst",
                approved_by="owner" if status == "Approved" else None,
            )
            requirement.rice_score = calculate_rice(requirement)
            db.add(requirement)
            requirements.append(requirement)
        db.flush()
        for requirement in requirements:
            db.add(RequirementVersion(
                requirement_id=requirement.id, version_number=1,
                snapshot_json=json.dumps(requirement_snapshot(requirement)),
                change_summary="Initial requirement", created_by="analyst",
            ))
            if requirement.id <= 5:
                data = generate_story(requirement)
                story = UserStory(
                    requirement_id=requirement.id, title=data["title"], persona=data["persona"],
                    goal=data["goal"], benefit=data["benefit"],
                    acceptance_criteria_json=json.dumps(data["acceptance_criteria"]),
                    definition_of_done_json=json.dumps(data["definition_of_done"]),
                    story_points=max(2, round(requirement.complexity)),
                    sprint="Sprint 1" if requirement.id <= 2 else "Backlog",
                )
                db.add(story)
                db.flush()
                db.add(TraceLink(
                    business_goal="Reduce requirement documentation time by 70%",
                    requirement_id=requirement.id, user_story_id=story.id,
                    task=f"TASK-{requirement.id:03d}",
                    test_case=f"TC-{requirement.id:03d}" if requirement.id <= 4 else "Not linked",
                    coverage_status="Covered" if requirement.id <= 4 else "Partial",
                ))
        db.add(ChangeRequest(
            requirement_id=requirements[1].id, summary="Add delegated approval path",
            reason="Product Owner may be unavailable during release freeze.", impact="Medium",
            requested_by="analyst",
        ))
        db.commit()


seed_database()

app = FastAPI(title="ReqFlow AI API", version="1.0.0", description="AI-powered requirement lifecycle management")
app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "http://localhost:5173,http://127.0.0.1:5173").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

REQUEST_COUNT = Counter("reqflow_http_requests_total", "HTTP requests", ["method", "path", "status"])
REQUEST_LATENCY = Histogram("reqflow_http_request_duration_seconds", "HTTP request latency", ["method", "path"])


@app.middleware("http")
async def observe_requests(request: Request, call_next):
    started = time.perf_counter()
    response = await call_next(request)
    elapsed = time.perf_counter() - started
    path = request.url.path
    REQUEST_COUNT.labels(request.method, path, response.status_code).inc()
    REQUEST_LATENCY.labels(request.method, path).observe(elapsed)
    response.headers["X-Response-Time-Ms"] = f"{elapsed * 1000:.2f}"
    return response


@app.get("/health")
def health() -> dict[str, Any]:
    return {"status": "healthy", "service": "reqflow-api", "ai_provider": "openai" if os.getenv("OPENAI_API_KEY") else "local"}


@app.get("/metrics")
def metrics() -> Response:
    return Response(generate_latest(), media_type=CONTENT_TYPE_LATEST)


@app.post("/api/auth/login")
def login(payload: LoginPayload, db: Session = Depends(get_db)) -> dict[str, Any]:
    user = db.scalar(select(User).where(User.username == payload.username))
    if not user or not user.active or not verify_password(payload.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid username or password")
    audit(db, user, "LOGIN", "user", user.id)
    db.commit()
    return {"access_token": encode_token(user), "token_type": "bearer", "user": {"username": user.username, "full_name": user.full_name, "role": user.role}}


@app.get("/api/auth/me")
def me(user: User = Depends(get_current_user)) -> dict[str, Any]:
    return {"username": user.username, "full_name": user.full_name, "role": user.role, "permissions": sorted(ROLE_PERMISSIONS.get(user.role, set()))}


@app.get("/api/dashboard/summary")
def dashboard_summary(db: Session = Depends(get_db), _: User = Depends(require("read"))) -> dict[str, Any]:
    requirements = list(db.scalars(select(Requirement)).all())
    stories = db.scalar(select(func.count(UserStory.id))) or 0
    trace_total = db.scalar(select(func.count(TraceLink.id))) or 0
    trace_covered = db.scalar(select(func.count(TraceLink.id)).where(TraceLink.coverage_status == "Covered")) or 0
    pending_changes = db.scalar(select(func.count(ChangeRequest.id)).where(ChangeRequest.status == "Pending")) or 0
    status_counts: dict[str, int] = {}
    priority_counts: dict[str, int] = {}
    for requirement in requirements:
        status_counts[requirement.status] = status_counts.get(requirement.status, 0) + 1
        priority_counts[requirement.priority] = priority_counts.get(requirement.priority, 0) + 1
    recent = sorted(requirements, key=lambda item: item.updated_at, reverse=True)[:5]
    return {
        "cards": {
            "requirements": len(requirements),
            "approved": status_counts.get("Approved", 0),
            "user_stories": stories,
            "traceability": round((trace_covered / trace_total * 100), 1) if trace_total else 0,
            "pending_changes": pending_changes,
            "quality_score": round(sum(local_analysis(item.description)["quality_score"] for item in requirements) / max(len(requirements), 1)),
        },
        "status_counts": status_counts,
        "priority_counts": priority_counts,
        "recent_requirements": [requirement_dict(item) for item in recent],
        "activity": [
            {"actor": log.actor, "action": log.action, "entity": f"{log.entity_type} {log.entity_id}", "created_at": log.created_at.isoformat()}
            for log in db.scalars(select(AuditLog).order_by(AuditLog.created_at.desc()).limit(6)).all()
        ],
    }


@app.get("/api/requirements")
def list_requirements(
    search: str = "", status: str = "", priority: str = "", category: str = "",
    db: Session = Depends(get_db), _: User = Depends(require("read")),
) -> list[dict[str, Any]]:
    statement = select(Requirement).order_by(Requirement.updated_at.desc())
    if search:
        pattern = f"%{search}%"
        statement = statement.where((Requirement.title.ilike(pattern)) | (Requirement.code.ilike(pattern)) | (Requirement.description.ilike(pattern)))
    if status:
        statement = statement.where(Requirement.status == status)
    if priority:
        statement = statement.where(Requirement.priority == priority)
    if category:
        statement = statement.where(Requirement.category == category)
    return [requirement_dict(item) for item in db.scalars(statement).all()]


@app.post("/api/requirements", status_code=201)
def create_requirement(payload: RequirementCreate, db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    requirement = Requirement(code=next_requirement_code(db), created_by=user.username, **payload.model_dump())
    requirement.rice_score = calculate_rice(requirement)
    db.add(requirement)
    db.flush()
    db.add(RequirementVersion(
        requirement_id=requirement.id, version_number=1,
        snapshot_json=json.dumps(requirement_snapshot(requirement)),
        change_summary="Initial requirement", created_by=user.username,
    ))
    audit(db, user, "CREATE", "requirement", requirement.code, requirement.title)
    db.commit()
    db.refresh(requirement)
    return requirement_dict(requirement)


@app.post("/api/requirements/import", status_code=201)
async def import_requirements(file: UploadFile = File(...), db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    raw = await file.read()
    try:
        rows = list(csv.DictReader(io.StringIO(raw.decode("utf-8-sig"))))
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=400, detail="CSV must be UTF-8 encoded") from exc
    created = []
    for row in rows:
        title = (row.get("title") or "").strip()
        description = (row.get("description") or "").strip()
        if not title or not description:
            continue
        requirement = Requirement(
            code=next_requirement_code(db), title=title, description=description,
            requirement_type=row.get("requirement_type") or "functional",
            category=row.get("category") or "General", priority=row.get("priority") or "Should",
            status=row.get("status") or "Draft", owner=row.get("owner") or user.full_name,
            created_by=user.username,
        )
        requirement.rice_score = calculate_rice(requirement)
        db.add(requirement)
        db.flush()
        db.add(RequirementVersion(requirement_id=requirement.id, version_number=1, snapshot_json=json.dumps(requirement_snapshot(requirement)), change_summary="Imported from CSV", created_by=user.username))
        created.append(requirement.code)
    audit(db, user, "IMPORT", "requirement", "bulk", f"{len(created)} requirements from {file.filename}")
    db.commit()
    return {"created": len(created), "codes": created}


@app.get("/api/requirements/{requirement_id}")
def get_requirement(requirement_id: int, db: Session = Depends(get_db), _: User = Depends(require("read"))) -> dict[str, Any]:
    requirement = db.get(Requirement, requirement_id)
    if not requirement:
        raise HTTPException(status_code=404, detail="Requirement not found")
    result = requirement_dict(requirement)
    result["analysis"] = local_analysis(requirement.description)
    result["stories"] = [story_dict(story, requirement.code) for story in db.scalars(select(UserStory).where(UserStory.requirement_id == requirement.id)).all()]
    return result


@app.put("/api/requirements/{requirement_id}")
def update_requirement(requirement_id: int, payload: RequirementUpdate, db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    requirement = db.get(Requirement, requirement_id)
    if not requirement:
        raise HTTPException(status_code=404, detail="Requirement not found")
    data = payload.model_dump(exclude_none=True)
    change_summary = data.pop("change_summary", "Requirement updated")
    for key, value in data.items():
        setattr(requirement, key, value)
    requirement.version += 1
    requirement.updated_at = utcnow()
    requirement.rice_score = calculate_rice(requirement)
    db.add(RequirementVersion(requirement_id=requirement.id, version_number=requirement.version, snapshot_json=json.dumps(requirement_snapshot(requirement)), change_summary=change_summary, created_by=user.username))
    audit(db, user, "UPDATE", "requirement", requirement.code, change_summary)
    db.commit()
    return requirement_dict(requirement)


@app.get("/api/requirements/{requirement_id}/history")
def requirement_history(requirement_id: int, db: Session = Depends(get_db), _: User = Depends(require("read"))) -> list[dict[str, Any]]:
    versions = db.scalars(select(RequirementVersion).where(RequirementVersion.requirement_id == requirement_id).order_by(RequirementVersion.version_number.desc())).all()
    return [{"id": item.id, "version": item.version_number, "snapshot": json.loads(item.snapshot_json), "change_summary": item.change_summary, "created_by": item.created_by, "created_at": item.created_at.isoformat()} for item in versions]


@app.post("/api/requirements/{requirement_id}/analyze")
def analyze_requirement(requirement_id: int, db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    requirement = db.get(Requirement, requirement_id)
    if not requirement:
        raise HTTPException(status_code=404, detail="Requirement not found")
    analysis = local_analysis(requirement.description)
    ai_raw = openai_answer(
        "You are a senior business analyst. Analyze the requirement and return concise JSON with keys: classification, category, functional_requirements, non_functional_requirements, stakeholders, ambiguities, missing_requirements, quality_score.",
        requirement.description,
    )
    if ai_raw:
        try:
            cleaned = re.sub(r"^```(?:json)?|```$", "", ai_raw.strip(), flags=re.MULTILINE).strip()
            analysis = {**analysis, **json.loads(cleaned), "provider": "openai", "model": OPENAI_MODEL}
        except json.JSONDecodeError:
            analysis["ai_summary"] = ai_raw
            analysis["provider"] = "openai"
            analysis["model"] = OPENAI_MODEL
    requirement.requirement_type = analysis.get("classification", requirement.requirement_type)
    requirement.category = analysis.get("category", requirement.category)
    audit(db, user, "ANALYZE", "requirement", requirement.code, f"Provider: {analysis['provider']}")
    db.commit()
    return analysis


@app.post("/api/requirements/{requirement_id}/user-stories", status_code=201)
def create_user_story(requirement_id: int, db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    requirement = db.get(Requirement, requirement_id)
    if not requirement:
        raise HTTPException(status_code=404, detail="Requirement not found")
    data = generate_story(requirement)
    story = UserStory(
        requirement_id=requirement.id, title=data["title"], persona=data["persona"], goal=data["goal"], benefit=data["benefit"],
        acceptance_criteria_json=json.dumps(data["acceptance_criteria"]), definition_of_done_json=json.dumps(data["definition_of_done"]),
        story_points=max(2, round(requirement.complexity)),
    )
    db.add(story)
    db.flush()
    db.add(TraceLink(business_goal="Improve requirement delivery quality", requirement_id=requirement.id, user_story_id=story.id, coverage_status="Partial"))
    audit(db, user, "GENERATE", "user_story", story.id, requirement.code)
    db.commit()
    return story_dict(story, requirement.code)


@app.get("/api/user-stories")
def list_user_stories(db: Session = Depends(get_db), _: User = Depends(require("read"))) -> list[dict[str, Any]]:
    requirements = {item.id: item.code for item in db.scalars(select(Requirement)).all()}
    return [story_dict(story, requirements.get(story.requirement_id, "")) for story in db.scalars(select(UserStory).order_by(UserStory.id.desc())).all()]


@app.post("/api/prioritization/recalculate")
def recalculate_priorities(db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    requirements = list(db.scalars(select(Requirement)).all())
    for requirement in requirements:
        requirement.rice_score = calculate_rice(requirement)
    audit(db, user, "RECALCULATE", "prioritization", "all", f"{len(requirements)} items")
    db.commit()
    return {"updated": len(requirements)}


@app.get("/api/prioritization/matrix")
def prioritization_matrix(db: Session = Depends(get_db), _: User = Depends(require("read"))) -> dict[str, Any]:
    requirements = list(db.scalars(select(Requirement).order_by(Requirement.rice_score.desc())).all())
    quadrants = {"Quick wins": [], "Strategic": [], "Fill-ins": [], "Defer": []}
    for item in requirements:
        if item.business_value >= 7 and item.complexity <= 5:
            quadrant = "Quick wins"
        elif item.business_value >= 7:
            quadrant = "Strategic"
        elif item.complexity <= 5:
            quadrant = "Fill-ins"
        else:
            quadrant = "Defer"
        quadrants[quadrant].append(requirement_dict(item))
    return {"ranked": [requirement_dict(item) for item in requirements], "quadrants": quadrants}


def build_document_content(document_type: str, requirements: list[Requirement]) -> str:
    heading = {"BRD": "Business Requirements Document", "SRS": "Software Requirements Specification", "BACKLOG": "Product Backlog"}[document_type]
    sections = [heading, "", "1. Purpose", "ReqFlow AI centralizes and governs the complete requirement lifecycle.", "", "2. Scope"]
    if document_type == "BRD":
        sections += ["Business objectives, stakeholder needs, priorities, and measurable outcomes.", "", "3. Business Requirements"]
    elif document_type == "SRS":
        sections += ["Functional behavior, quality attributes, data, security, and integration expectations.", "", "3. System Requirements"]
    else:
        sections += ["Prioritized delivery items ready for refinement and sprint planning.", "", "3. Backlog Items"]
    for requirement in requirements:
        sections += [f"{requirement.code} - {requirement.title}", f"Priority: {requirement.priority} | Status: {requirement.status} | Owner: {requirement.owner}", requirement.description, ""]
    sections += ["4. Governance", "All changes require version history, traceability review, and role-based approval."]
    return "\n".join(sections)


@app.post("/api/documents/generate", status_code=201)
def generate_document(payload: DocumentGeneratePayload, db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    statement = select(Requirement)
    if payload.requirement_ids:
        statement = statement.where(Requirement.id.in_(payload.requirement_ids))
    requirements = list(db.scalars(statement.order_by(Requirement.code)).all())
    if not requirements:
        raise HTTPException(status_code=400, detail="No requirements selected")
    name = payload.name or f"ReqFlow AI {payload.document_type} v1.0"
    document = GeneratedDocument(name=name, document_type=payload.document_type, content=build_document_content(payload.document_type, requirements), requirement_count=len(requirements), generated_by=user.username)
    db.add(document)
    db.flush()
    audit(db, user, "GENERATE", "document", document.id, payload.document_type)
    db.commit()
    return {"id": document.id, "name": document.name, "document_type": document.document_type, "status": document.status, "requirement_count": document.requirement_count, "generated_by": document.generated_by, "created_at": document.created_at.isoformat()}


@app.get("/api/documents")
def list_documents(db: Session = Depends(get_db), _: User = Depends(require("read"))) -> list[dict[str, Any]]:
    documents = db.scalars(select(GeneratedDocument).order_by(GeneratedDocument.created_at.desc())).all()
    return [{"id": item.id, "name": item.name, "document_type": item.document_type, "status": item.status, "requirement_count": item.requirement_count, "generated_by": item.generated_by, "created_at": item.created_at.isoformat()} for item in documents]


@app.get("/api/documents/{document_id}/export")
def export_document(document_id: int, format: str = Query(pattern="^(pdf|docx)$"), db: Session = Depends(get_db), user: User = Depends(require("export"))):
    document = db.get(GeneratedDocument, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    buffer = io.BytesIO()
    if format == "pdf":
        styles = getSampleStyleSheet()
        story = []
        for index, line in enumerate(document.content.splitlines()):
            if not line:
                story.append(Spacer(1, 8))
            else:
                style = styles["Heading1"] if index == 0 else styles["Heading3"] if re.match(r"^\d+\.", line) else styles["BodyText"]
                story.append(Paragraph(line.replace("&", "&amp;"), style))
        SimpleDocTemplate(buffer, pagesize=A4, title=document.name).build(story)
        media_type = "application/pdf"
    else:
        docx = DocxDocument()
        for index, line in enumerate(document.content.splitlines()):
            if index == 0:
                docx.add_heading(line, 0)
            elif re.match(r"^\d+\.", line):
                docx.add_heading(line, level=1)
            else:
                docx.add_paragraph(line)
        docx.save(buffer)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    buffer.seek(0)
    audit(db, user, "EXPORT", "document", document.id, format.upper())
    db.commit()
    filename = re.sub(r"[^a-zA-Z0-9_-]+", "-", document.name).strip("-").lower()
    return StreamingResponse(buffer, media_type=media_type, headers={"Content-Disposition": f'attachment; filename="{filename}.{format}"'})


@app.post("/api/process-models/generate")
def generate_process_model(payload: ProcessGeneratePayload, db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    requirement = db.get(Requirement, payload.requirement_id) if payload.requirement_id else db.scalar(select(Requirement).order_by(Requirement.id))
    if not requirement:
        raise HTTPException(status_code=404, detail="Requirement not found")
    if payload.model_type == "USE_CASE":
        result = {"title": f"Use case: {requirement.title}", "actor": requirement.owner, "preconditions": ["User is authenticated", "Requirement is accessible"], "main_flow": ["User opens the workspace", f"User initiates {requirement.title.lower()}", "System validates the request", "System stores the outcome", "System records an audit event"], "alternate_flows": ["Validation fails and actionable errors are returned", "Authorization fails and access is denied"]}
    elif payload.model_type == "BPMN":
        result = {"title": f"Process: {requirement.title}", "nodes": [{"id": "start", "type": "event", "label": "Request received"}, {"id": "review", "type": "task", "label": "Analyze requirement"}, {"id": "decision", "type": "gateway", "label": "Complete?"}, {"id": "approve", "type": "task", "label": "Approve & baseline"}, {"id": "revise", "type": "task", "label": "Request changes"}, {"id": "end", "type": "event", "label": "Published"}], "edges": [["start", "review"], ["review", "decision"], ["decision", "approve"], ["decision", "revise"], ["revise", "review"], ["approve", "end"]]}
    else:
        result = {"title": "ReqFlow domain model", "entities": [{"name": "Requirement", "fields": ["id PK", "code", "title", "status", "version"]}, {"name": "UserStory", "fields": ["id PK", "requirement_id FK", "persona", "goal"]}, {"name": "TraceLink", "fields": ["id PK", "requirement_id FK", "user_story_id FK", "test_case"]}, {"name": "ChangeRequest", "fields": ["id PK", "requirement_id FK", "status", "impact"]}], "relationships": ["Requirement 1:N UserStory", "Requirement 1:N ChangeRequest", "Requirement 1:N TraceLink", "UserStory 1:N TraceLink"]}
    audit(db, user, "GENERATE", "process_model", requirement.code, payload.model_type)
    db.commit()
    return {"model_type": payload.model_type, "requirement": requirement_dict(requirement), "model": result}


@app.get("/api/traceability")
def traceability(db: Session = Depends(get_db), _: User = Depends(require("read"))) -> dict[str, Any]:
    requirements = {item.id: item for item in db.scalars(select(Requirement)).all()}
    stories = {item.id: item for item in db.scalars(select(UserStory)).all()}
    rows = []
    for link in db.scalars(select(TraceLink).order_by(TraceLink.id)).all():
        requirement = requirements.get(link.requirement_id)
        story = stories.get(link.user_story_id) if link.user_story_id else None
        rows.append({"id": link.id, "business_goal": link.business_goal, "requirement_code": requirement.code if requirement else "", "requirement_title": requirement.title if requirement else "", "user_story": story.title if story else "Not linked", "task": link.task, "test_case": link.test_case, "coverage_status": link.coverage_status})
    covered = sum(1 for row in rows if row["coverage_status"] == "Covered")
    return {"rows": rows, "coverage": round(covered / len(rows) * 100, 1) if rows else 0, "gaps": [row for row in rows if row["coverage_status"] != "Covered"]}


@app.post("/api/traceability/links", status_code=201)
def create_trace(payload: TraceCreatePayload, db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    if not db.get(Requirement, payload.requirement_id):
        raise HTTPException(status_code=404, detail="Requirement not found")
    link = TraceLink(**payload.model_dump())
    db.add(link)
    db.flush()
    audit(db, user, "CREATE", "trace_link", link.id)
    db.commit()
    return {"id": link.id, **payload.model_dump()}


@app.get("/api/change-requests")
def list_changes(db: Session = Depends(get_db), _: User = Depends(require("read"))) -> list[dict[str, Any]]:
    requirements = {item.id: item for item in db.scalars(select(Requirement)).all()}
    changes = db.scalars(select(ChangeRequest).order_by(ChangeRequest.created_at.desc())).all()
    return [{"id": item.id, "requirement_id": item.requirement_id, "requirement_code": requirements[item.requirement_id].code, "requirement_title": requirements[item.requirement_id].title, "summary": item.summary, "reason": item.reason, "impact": item.impact, "status": item.status, "requested_by": item.requested_by, "reviewed_by": item.reviewed_by, "created_at": item.created_at.isoformat()} for item in changes]


@app.post("/api/change-requests", status_code=201)
def create_change(payload: ChangeCreatePayload, db: Session = Depends(get_db), user: User = Depends(require("write"))) -> dict[str, Any]:
    if not db.get(Requirement, payload.requirement_id):
        raise HTTPException(status_code=404, detail="Requirement not found")
    change = ChangeRequest(**payload.model_dump(), requested_by=user.username)
    db.add(change)
    db.flush()
    audit(db, user, "REQUEST_CHANGE", "requirement", payload.requirement_id, payload.summary)
    db.commit()
    return {"id": change.id, "status": change.status}


@app.post("/api/change-requests/{change_id}/decision")
def decide_change(change_id: int, payload: ChangeDecisionPayload, db: Session = Depends(get_db), user: User = Depends(require("approve"))) -> dict[str, Any]:
    change = db.get(ChangeRequest, change_id)
    if not change:
        raise HTTPException(status_code=404, detail="Change request not found")
    change.status = payload.decision
    change.reviewed_by = user.username
    change.reviewed_at = utcnow()
    requirement = db.get(Requirement, change.requirement_id)
    if requirement and payload.decision == "Approved":
        requirement.status = "Approved"
        requirement.approved_by = user.username
    audit(db, user, payload.decision.upper(), "change_request", change.id, change.summary)
    db.commit()
    return {"id": change.id, "status": change.status, "reviewed_by": change.reviewed_by}


@app.post("/api/assistant/chat")
def assistant_chat(payload: AssistantPayload, db: Session = Depends(get_db), user: User = Depends(require("read"))) -> dict[str, Any]:
    requirement = db.get(Requirement, payload.requirement_id) if payload.requirement_id else None
    context = requirement.description if requirement else "\n".join(f"{item.code}: {item.title} - {item.description}" for item in db.scalars(select(Requirement).limit(12)).all())
    answer = openai_answer("You are ReqFlow AI, a senior Business Analyst assistant. Answer concisely, identify assumptions, and propose testable requirement language.", f"Context:\n{context}\n\nQuestion: {payload.question}")
    provider = "openai" if answer else "local"
    if not answer:
        analysis = local_analysis(context)
        lowered = payload.question.lower()
        if "missing" in lowered or "gap" in lowered:
            answer = "Review these gaps: " + " ".join(analysis["missing_requirements"])
        elif "stakeholder" in lowered:
            answer = "Likely stakeholders: " + ", ".join(analysis["stakeholders"]) + ". Validate decision authority and notification needs with each group."
        elif "acceptance" in lowered:
            answer = "Use observable outcomes: successful persistence, validation without data loss, RBAC enforcement, audit evidence, and a measurable response-time threshold."
        else:
            answer = f"The current scope contains {db.scalar(select(func.count(Requirement.id))) or 0} requirements. Focus next on measurable acceptance criteria, failure behavior, authorization, and complete links to test cases."
    audit(db, user, "ASK_AI", "assistant", requirement.code if requirement else "workspace", payload.question[:120])
    db.commit()
    return {"answer": answer, "provider": provider, "model": OPENAI_MODEL if provider == "openai" else "ReqFlow local analyst", "suggested_actions": ["Review ambiguity", "Add acceptance criteria", "Check traceability gaps"]}


@app.get("/api/assistant/suggestions")
def assistant_suggestions(db: Session = Depends(get_db), _: User = Depends(require("read"))) -> list[dict[str, Any]]:
    suggestions = []
    for requirement in db.scalars(select(Requirement).order_by(Requirement.updated_at.desc())).all():
        analysis = local_analysis(requirement.description)
        if analysis["ambiguities"] or requirement.status == "Draft":
            suggestions.append({"requirement_id": requirement.id, "code": requirement.code, "title": requirement.title, "quality_score": analysis["quality_score"], "issue": analysis["ambiguities"][0] if analysis["ambiguities"] else "Draft requirement is awaiting refinement.", "action": analysis["missing_requirements"][0]})
    return suggestions[:6]


@app.get("/api/audit-logs")
def audit_logs(limit: int = Query(default=50, ge=1, le=200), db: Session = Depends(get_db), _: User = Depends(require("read"))) -> list[dict[str, Any]]:
    logs = db.scalars(select(AuditLog).order_by(AuditLog.created_at.desc()).limit(limit)).all()
    return [{"id": item.id, "actor": item.actor, "action": item.action, "entity_type": item.entity_type, "entity_id": item.entity_id, "detail": item.detail, "created_at": item.created_at.isoformat()} for item in logs]


@app.get("/api/admin/users")
def list_users(db: Session = Depends(get_db), _: User = Depends(require("manage_users"))) -> list[dict[str, Any]]:
    return [{"id": user.id, "username": user.username, "full_name": user.full_name, "role": user.role, "active": user.active} for user in db.scalars(select(User).order_by(User.id)).all()]
